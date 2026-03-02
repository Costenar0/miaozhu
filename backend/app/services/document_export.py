"""Word / PDF 文档导出

软著申请格式要求：
- 操作说明书：不少于 60 页，每页不少于 30 行
- 源程序代码：不少于 60 页，每页不少于 50 行
"""

import io
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import mistletoe
    from mistletoe import Document as MdDocument
    from mistletoe.block_token import Heading, List, CodeFence, Paragraph as MdParagraph, Table
    from mistletoe.span_token import RawText, Strong
    _MISTLETOE_AVAILABLE = True
except ImportError:
    _MISTLETOE_AVAILABLE = False

from app.models.application import Application
from app.models.generation import GenerationSection


# ==================== XML 安全辅助 ====================

# XML 1.0 合法字符：#x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
_XML_ILLEGAL = re.compile(
    r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)


def _clean(text: str) -> str:
    """移除 XML 1.0 不允许的控制字符（NULL 字节、\x01-\x08 等），防止 python-docx 写入失败。"""
    if not text:
        return text
    return _XML_ILLEGAL.sub("", text)


# ==================== 常量 ====================

# 操作说明书：宋体 小四(12pt)，固定行距 22pt → ~31 行/页，满足 ≥30 行要求
_MANUAL_FONT_SIZE = Pt(12)
_MANUAL_LINE_SPACING = Pt(22)

# 源程序代码：Times New Roman 五号(10.5pt)，固定行距 13pt → ~53 行/页，满足 ≥50 行要求
_CODE_FONT = "Times New Roman"
_CODE_FONT_SIZE = Pt(10.5)
_CODE_LINE_SPACING = Pt(13)

# 中文章节编号
_CN_CHAPTER_NUMS = [
    "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
    "十一", "十二", "十三", "十四", "十五",
]


class _HeadingCounter:
    """章节内子标题编号计数器（1.1, 1.2, 1.1.1, 1.1.2 ...）"""

    def __init__(self, chapter_num: int):
        self.chapter = chapter_num
        self.h2 = 0
        self.h3 = 0

    def next_h2(self, title: str) -> str:
        self.h2 += 1
        self.h3 = 0
        return f"{self.chapter}.{self.h2}  {title}"

    def next_h3(self, title: str) -> str:
        self.h3 += 1
        return f"{self.chapter}.{self.h2}.{self.h3}  {title}"


def _strip_numbering(title: str) -> str:
    """移除标题中的编号前缀（数字编号或中文编号）

    Examples:
        "1.1 功能说明" → "功能说明"
        "1.1.1 操作步骤" → "操作步骤"
        "一、引言" → "引言"
        "（一）概述" → "概述"
    """
    # 移除数字编号：1. / 1.1 / 1.1.1 / 1、等（支持空格或无空格）
    title = re.sub(r'^\d+(\.\d+)*[、.．、]?\s*', '', title)
    # 移除中文编号：一、 / （一） 等
    title = re.sub(r'^[（(]?[一二三四五六七八九十百千]+[）)、]\s*', '', title)
    # 移除括号编号：(1) / （1）等
    title = re.sub(r'^[（(]\d+[）)]\s*', '', title)
    return title.strip()


def _extract_headings(content: str) -> list[tuple[int, str]]:
    """从 Markdown 内容中提取标题层级和文本（自动去除编号）"""
    headings: list[tuple[int, str]] = []

    if _MISTLETOE_AVAILABLE:
        # 使用 mistletoe 精确解析
        try:
            md_doc = MdDocument(content)
            for token in md_doc.children:
                if isinstance(token, Heading):
                    title = _render_span_tokens(token.children, paragraph=None, is_manual=False)
                    title = _strip_numbering(title)
                    headings.append((token.level, title))
            return headings
        except Exception:
            pass  # 失败则回退到正则解析

    # 回退到逐行正则解析
    in_code_block = False
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            if title:
                # 移除标题中已有的编号
                title = _strip_numbering(title)
                headings.append((hashes, title))
    return headings


def _merge_manual_chapters(
    sections: list[GenerationSection],
) -> list[tuple[str, str]]:
    """将 manual sections 合并为章节列表，功能详述（上/下）合并为一个。

    Returns: [(title, content), ...]
    """
    chapters: list[tuple[str, str]] = []
    func_parts: list[str] = []

    for s in sections:
        if not s.content:
            continue
        if s.section_key.startswith("manual_functions"):
            func_parts.append(s.content)
        else:
            if func_parts:
                chapters.append(("功能详述", "\n\n".join(func_parts)))
                func_parts = []
            chapters.append((s.title, s.content))

    if func_parts:
        chapters.append(("功能详述", "\n\n".join(func_parts)))

    return chapters


def _build_toc_entries(
    chapters: list[tuple[str, str]],
) -> list[tuple[int, str]]:
    """根据章节列表构建 TOC 条目 [(level, numbered_title), ...]"""
    entries: list[tuple[int, str]] = []
    for i, (title, content) in enumerate(chapters):
        cn = _CN_CHAPTER_NUMS[i] if i < len(_CN_CHAPTER_NUMS) else str(i + 1)
        entries.append((1, f"{cn}、{title}"))

        chapter_num = i + 1
        h2_count = 0
        h3_count = 0
        for h_level, h_title in _extract_headings(content):
            if h_level == 1:
                h2_count += 1
                h3_count = 0
                entries.append((2, f"{chapter_num}.{h2_count}  {h_title}"))
            elif h_level >= 2 and h2_count > 0:  # 只有在有二级标题时才添加三级标题
                h3_count += 1
                entries.append((3, f"{chapter_num}.{h2_count}.{h3_count}  {h_title}"))

    return entries


# ==================== Word 通用辅助 ====================

def _setup_page(doc: Document) -> None:
    """设置 A4 页面和页边距"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def _add_page_numbers(doc: Document) -> None:
    """给所有节添加页码（页脚居中）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run1._r.append(fld_begin)

        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run2._r.append(instr)

        run3 = p.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_end)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(_clean(text), level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        _set_cn_font(run, "黑体")

    # 禁用自动编号（移除段落的编号属性）
    pPr = heading._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)


def _set_cn_font(run, font_name: str = "宋体") -> None:
    """设置中文字体（eastAsia）和西文字体（Times New Roman）"""
    run.font.name = "Times New Roman"
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)


def _add_manual_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """添加操作说明书正文段落，宋体 12pt，固定行距 22pt"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _MANUAL_LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if bold:
        run = p.add_run(_clean(text))
        run.bold = True
        _set_cn_font(run)
        run.font.size = _MANUAL_FONT_SIZE
    else:
        _add_inline_formatting(p, text, is_manual=True)


def _add_inline_formatting(paragraph, text: str, is_manual: bool = False) -> None:
    """处理行内 **粗体** 格式"""
    parts = re.split(r"(\*\*.*?\*\*)", _clean(text))
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        if is_manual:
            _set_cn_font(run)
            run.font.size = _MANUAL_FONT_SIZE


def _parse_table_row(line: str) -> list[str]:
    """解析 Markdown 表格行，返回单元格列表"""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _is_separator_row(line: str) -> bool:
    """判断是否是表格分隔行（如 | --- | --- |）"""
    return bool(re.match(r"^\|?[\s\-:]+(\|[\s\-:]+)+\|?\s*$", line))


def _add_word_table(doc: Document, table_lines: list[str]) -> None:
    """将 Markdown 表格行渲染为 Word 表格"""
    rows_data: list[list[str]] = []
    for line in table_lines:
        if _is_separator_row(line):
            continue
        rows_data.append(_parse_table_row(line))

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    # 表格占满页面宽度（5000 = 100%）
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # 细边框 0.5pt 灰色
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border_el = OxmlElement(f'w:{edge}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:color'), '999999')
        border_el.set(qn('w:space'), '0')
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = _MANUAL_LINE_SPACING
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            _add_inline_formatting(p, cell_text, is_manual=True)
            # 首行加粗作为表头
            if i == 0:
                for run in p.runs:
                    run.bold = True


def _add_list_item(doc: Document, text: str, style: str) -> None:
    """添加列表项，支持行内 **粗体**"""
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = _MANUAL_LINE_SPACING
    _add_inline_formatting(p, text, is_manual=True)


def _render_span_tokens(tokens: list, paragraph=None, is_manual: bool = True):
    """递归渲染 span tokens (inline elements) 到 Word paragraph

    Args:
        tokens: mistletoe span tokens
        paragraph: Word paragraph object
        is_manual: 是否使用操作说明书格式

    Returns:
        如果 paragraph 为 None，返回纯文本；否则在 paragraph 中添加 runs
    """
    text_parts = []

    for token in tokens:
        if hasattr(token, 'children'):
            # Recursively handle nested tokens (like Strong)
            child_text = _render_span_tokens(token.children, paragraph=None, is_manual=is_manual)

            if paragraph is not None:
                run = paragraph.add_run(_clean(child_text))
                if isinstance(token, Strong):
                    run.bold = True
                if is_manual:
                    _set_cn_font(run)
                    run.font.size = _MANUAL_FONT_SIZE
            else:
                text_parts.append(child_text)
        else:
            # Leaf token (RawText)
            content = token.content if hasattr(token, 'content') else str(token)
            if paragraph is not None:
                run = paragraph.add_run(_clean(content))
                if is_manual:
                    _set_cn_font(run)
                    run.font.size = _MANUAL_FONT_SIZE
            else:
                text_parts.append(content)

    if paragraph is None:
        return ''.join(text_parts)
    return ''


_DIAGRAM_RENDER_LANGS = {"mermaid", "plantuml", "d2", "html"}


def _render_diagram_png(lang: str, code: str) -> bytes | None:
    """将 mermaid/plantuml/d2/html 代码渲染为 PNG，失败返回 None。"""
    try:
        from app.services.diagram_renderer import render_mermaid, render_plantuml, render_d2, render_html_screenshot
        if lang == "mermaid":
            return render_mermaid(code)
        elif lang == "plantuml":
            if not code.strip().startswith("@startuml"):
                code = f"@startuml\n{code.strip()}\n@enduml"
            return render_plantuml(code)
        elif lang == "d2":
            return render_d2(code)
        elif lang == "html":
            return render_html_screenshot(code)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning("Diagram render failed (%s): %s", lang, e)
    return None


def _add_diagram_to_doc(doc: Document, png: bytes, lang: str) -> None:
    """将 PNG 图片插入 Word，自动适配页面尺寸（最大宽 14cm，最大高 18cm）。"""
    # A4 内容区最大宽约 15.9cm，最大高约 25.7cm，预留标题和文字后最大高 18cm
    max_w_cm = 14.0
    max_h_cm = 18.0
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(png))
        w_px, h_px = img.size
        aspect = h_px / max(w_px, 1)
        if aspect * max_w_cm <= max_h_cm:
            doc.add_picture(io.BytesIO(png), width=Cm(max_w_cm))
        else:
            doc.add_picture(io.BytesIO(png), height=Cm(max_h_cm))
    except Exception:
        # PIL 不可用时的回退：plantuml 使用较小宽度避免溢出
        fallback_w = Cm(10) if lang == "plantuml" else Cm(max_w_cm)
        doc.add_picture(io.BytesIO(png), width=fallback_w)


def _process_md_tokens(
    doc: Document,
    tokens: list,
    counter: _HeadingCounter | None = None,
    skip_first_h1: str | None = None,
    render_diagrams: bool = False,
) -> None:
    """处理 mistletoe 解析的 tokens，生成 Word 内容

    Args:
        doc: Word document
        tokens: mistletoe block tokens
        counter: 章节编号计数器（提供时，# → Heading 2 (编号 N.1)，## → Heading 3 (编号 N.1.1)）
        skip_first_h1: 如果提供，则跳过第一个与此标题相同的一级标题（避免重复）
    """
    first_h1_skipped = skip_first_h1 is None  # 如果不需要跳过，直接标记为已跳过

    for token in tokens:
        if isinstance(token, Heading):
            # 标题：提取纯文本并去除编号
            title_text = _render_span_tokens(token.children, paragraph=None, is_manual=False)
            title_text = _strip_numbering(title_text)

            # 跳过第一个与章节标题相同的一级标题（避免重复）
            if not first_h1_skipped and token.level == 1 and skip_first_h1:
                if title_text.strip() == skip_first_h1.strip():
                    first_h1_skipped = True
                    continue

            if counter:
                if token.level == 1:
                    _add_heading(doc, counter.next_h2(title_text), level=2)
                elif token.level >= 2:
                    _add_heading(doc, counter.next_h3(title_text), level=3)
            else:
                level = min(2 + token.level, 9)
                _add_heading(doc, title_text, level=level)

        elif isinstance(token, List):
            # 列表
            is_ordered = token.start is not None
            # 有序列表：用 List Paragraph 保留原始编号文本，避免 Word 全局计数器跨章节累加
            num_counter = token.start if is_ordered else 0
            for item in token.children:
                # List item 的 children 是 block tokens (通常是 Paragraph)
                for child in item.children:
                    if isinstance(child, MdParagraph):
                        text = _render_span_tokens(child.children, paragraph=None, is_manual=True)
                        if is_ordered:
                            _add_list_item(doc, f"{num_counter}. {text}", "List Paragraph")
                            num_counter += 1
                        else:
                            _add_list_item(doc, text, "List Bullet")
                    else:
                        # 嵌套的复杂结构，递归处理
                        _process_md_tokens(doc, [child], counter=None, skip_first_h1=None, render_diagrams=render_diagrams)

        elif isinstance(token, CodeFence):
            # 代码块
            lang = getattr(token, "language", "").strip().lower()
            code_text = token.children[0].content if token.children else ""
            if render_diagrams and lang in _DIAGRAM_RENDER_LANGS:
                png = _render_diagram_png(lang, code_text)
                if png:
                    _add_diagram_to_doc(doc, png, lang)
                    doc.add_paragraph()  # 图片后空一行
            else:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(16)
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                pf.left_indent = Cm(1)
                run = p.add_run(_clean(code_text.rstrip()))
                run.font.name = _CODE_FONT
                run.font.size = Pt(9)

        elif isinstance(token, Table):
            # 表格：转换为 markdown 文本行格式（复用现有表格处理函数）
            table_lines = []
            for row in token.children:
                cells = []
                for cell in row.children:
                    cell_text = _render_span_tokens(cell.children, paragraph=None, is_manual=True)
                    cells.append(cell_text)
                table_lines.append("| " + " | ".join(cells) + " |")
            if table_lines:
                _add_word_table(doc, table_lines)

        elif isinstance(token, MdParagraph):
            # 段落
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = _MANUAL_LINE_SPACING
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            _render_span_tokens(token.children, paragraph=p, is_manual=True)


def _add_manual_content(
    doc: Document,
    content: str,
    base_heading_level: int = 3,
    counter: _HeadingCounter | None = None,
    skip_first_h1: str | None = None,
    render_diagrams: bool = False,
) -> None:
    """操作说明书内容：Markdown → docx，宋体 12pt，固定行距 22pt（~31 行/页）

    Args:
        counter: 提供时，# → Heading 2 (编号 N.1)，## → Heading 3 (编号 N.1.1)
        skip_first_h1: 如果提供，则跳过第一个与此标题相同的一级标题（避免章节标题重复）
    """
    if _MISTLETOE_AVAILABLE:
        # 使用 mistletoe 解析 markdown
        try:
            md_doc = MdDocument(content)
            _process_md_tokens(doc, md_doc.children, counter=counter, skip_first_h1=skip_first_h1, render_diagrams=render_diagrams)
            return
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning("Mistletoe parsing failed, falling back to regex: %s", e)

    # 回退到原有的逐行正则解析（兼容性）
    lines = content.split("\n")
    code_block = False
    code_block_lang = ""
    code_lines: list[str] = []
    table_lines: list[str] = []
    first_h1_skipped = skip_first_h1 is None  # 如果不需要跳过，直接标记为已跳过

    for line in lines:
        stripped = line.strip()

        # 如果正在收集表格行
        if table_lines:
            if stripped.startswith("|"):
                table_lines.append(stripped)
                continue
            else:
                _add_word_table(doc, table_lines)
                table_lines = []

        # 代码块
        if stripped.startswith("```"):
            if code_block:
                code_text = "\n".join(code_lines)
                if render_diagrams and code_block_lang in _DIAGRAM_RENDER_LANGS:
                    png = _render_diagram_png(code_block_lang, code_text)
                    if png:
                        _add_diagram_to_doc(doc, png, code_block_lang)
                        doc.add_paragraph()
                else:
                    p = doc.add_paragraph()
                    pf = p.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = Pt(16)
                    pf.space_before = Pt(4)
                    pf.space_after = Pt(4)
                    pf.left_indent = Cm(1)
                    run = p.add_run(_clean(code_text))
                    run.font.name = _CODE_FONT
                    run.font.size = Pt(9)
                code_lines = []
                code_block = False
                code_block_lang = ""
            else:
                code_block = True
                code_block_lang = stripped[3:].strip().lower()  # 捕获语言标签
            continue

        if code_block:
            code_lines.append(line)
            continue

        # Markdown 标题
        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped.lstrip("#").strip()
            if title_text:
                # 移除标题中已有的编号（如 "1.1 功能说明" → "功能说明"）
                title_text = _strip_numbering(title_text)

                # 跳过第一个与章节标题相同的一级标题（避免重复）
                if not first_h1_skipped and hashes == 1 and skip_first_h1:
                    if title_text.strip() == skip_first_h1.strip():
                        first_h1_skipped = True
                        continue

                if counter:
                    if hashes == 1:
                        _add_heading(doc, counter.next_h2(title_text), level=2)
                    elif hashes >= 2:
                        _add_heading(doc, counter.next_h3(title_text), level=3)
                else:
                    level = min(base_heading_level + hashes - 1, 9)
                    _add_heading(doc, title_text, level=level)
                continue

        # 表格开始
        if stripped.startswith("|"):
            table_lines.append(stripped)
            continue

        # 列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_list_item(doc, stripped[2:], "List Bullet")
        elif re.match(r"^\d+\.\s*", stripped):
            # 有序列表：保留原始编号文本，用 List Paragraph（仅缩进）
            # 不使用 "List Number" 样式，避免 Word 全局计数器跨章节累加到 300+
            _add_list_item(doc, stripped, "List Paragraph")
        elif stripped == "":
            continue
        else:
            _add_manual_paragraph(doc, stripped)

    # 未关闭的表格
    if table_lines:
        _add_word_table(doc, table_lines)

    # 未关闭的代码块
    if code_lines:
        code_text = "\n".join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(_clean(code_text))
        run.font.name = _CODE_FONT
        run.font.size = Pt(9)


def _add_source_code_content(doc: Document, content: str) -> None:
    """源程序代码内容：Courier New 10.5pt，固定行距 13pt（~53 行/页）

    源代码以纯文本渲染，每行一个段落，确保精确的行密度控制。
    """
    # 去除可能的 markdown 代码块标记
    text = content.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        # 去掉首尾的 ``` 行
        cleaned = []
        in_code = False
        for line in lines:
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            cleaned.append(line)
        text = "\n".join(cleaned) if cleaned else text

    for line in text.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = _CODE_LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(_clean(line))
        run.font.name = _CODE_FONT
        run.font.size = _CODE_FONT_SIZE


def _init_word_doc() -> Document:
    """初始化 Word 文档：A4 页面 + 默认字体"""
    doc = Document()
    _setup_page(doc)

    # 设置 Normal 样式字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), '宋体')

    # 禁用 Heading 1/2/3 样式的自动编号
    # 注意：保留编号属性的移除，但实际上 Heading 样式默认不带编号
    # 这里主要是为了确保在某些模板中加载的样式不会带编号
    for heading_level in [1, 2, 3]:
        try:
            heading_style = doc.styles[f'Heading {heading_level}']
            pPr = heading_style.element.get_or_add_pPr()
            # 移除编号属性
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
        except KeyError:
            # 样式不存在，跳过
            pass

    return doc


def _add_cover_page(doc: Document, app: Application, doc_title: str) -> None:
    """添加封面页

    格式：
    软件名称
    使用说明书
    (版本V1.0)

    申请人：XXX有限公司
    申请日期：YYYY年MM月DD日
    """
    # 软件名称（顶部留白 8cm）
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Cm(8)
    p1.paragraph_format.space_after = Pt(12)
    run1 = p1.add_run(_clean(app.software_name or "软件系统"))
    run1.font.size = Pt(22)
    run1.bold = True
    _set_cn_font(run1, "黑体")

    # 文档标题（使用说明书 / 源程序鉴别材料等）
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(_clean(doc_title))
    run2.font.size = Pt(18)
    run2.bold = True
    _set_cn_font(run2, "黑体")

    # 版本号
    version = app.software_version or "V1.0"
    if not version.startswith("V") and not version.startswith("v"):
        version = f"V{version}"
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Cm(4)
    run3 = p3.add_run(_clean(f"(版本{version})"))
    run3.font.size = Pt(14)
    _set_cn_font(run3, "宋体")

    # 申请人
    applicant = app.applicant_name or "申请单位"
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(8)
    run4 = p4.add_run(_clean(f"申请人：{applicant}"))
    run4.font.size = Pt(14)
    _set_cn_font(run4, "宋体")

    # 申请日期
    from datetime import datetime
    if app.completion_date:
        # completion_date 可能是 "2024-08-12" 或 "2024年8月12日" 格式
        date_str = app.completion_date
        if "-" in date_str:
            # 转换为中文格式
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d")
                date_str = f"{dt.year}年{dt.month}月{dt.day}日"
            except ValueError:
                # 如果格式不对，直接使用原值
                pass
    else:
        # 使用当前日期
        now = datetime.now()
        date_str = f"{now.year}年{now.month}月{now.day}日"

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(_clean(f"申请日期：{date_str}"))
    run5.font.size = Pt(14)
    _set_cn_font(run5, "宋体")

    # 添加分页符
    doc.add_page_break()


def _save_doc(doc: Document) -> io.BytesIO:
    """添加页码并保存到 BytesIO"""
    _add_page_numbers(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==================== Word 目录 ====================

def _add_word_toc(doc: Document, entries: list[tuple[int, str]]) -> None:
    """添加 Word 目录页（TOC 域代码 + 简化预填充 + 分节符）

    目录使用 Word TOC 域代码，dirty 标志让 Word/WPS 打开时自动刷新。
    简化预填充逻辑以提高性能。

    Args:
        entries: [(level, numbered_title), ...] — 所有目录条目
    """
    import logging
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_SECTION

    logger = logging.getLogger(__name__)
    logger.info("Generating TOC with %d entries", len(entries))

    # ── 目录标题 ──
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(40)
    heading.paragraph_format.space_after = Pt(24)
    run = heading.add_run("目  录")
    run.bold = True
    run.font.size = Pt(18)
    _set_cn_font(run, "黑体")

    # ── 配置 TOC 样式（缩进为 0，右对齐制表位 + 前导点） ──
    # 可用宽度：21cm - 2.54cm(左边距) - 2.54cm(右边距) = 15.92cm
    # 转换为 twips (1 cm = 567 twips): 15.92 * 567 ≈ 9027 twips
    toc_tab_pos = 9027

    toc_style_cfgs = [
        ("toc 1", Pt(12), True, "黑体"),
        ("toc 2", Pt(11), False, "宋体"),
        ("toc 3", Pt(10.5), False, "宋体"),
    ]
    for name, size, bold, cn_font in toc_style_cfgs:
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        pf = style.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(1)
        pf.left_indent = Cm(0)
        style.font.size = size
        style.font.bold = bold
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), cn_font)
        # 右对齐制表位 + dot leader（Word 更新域后用于对齐页码）
        tab_xml = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:leader'), 'dot')
        tab_el.set(qn('w:pos'), str(toc_tab_pos))
        tab_xml.append(tab_el)
        pf_xml = style.element.find(qn('w:pPr'))
        if pf_xml is None:
            pf_xml = OxmlElement('w:pPr')
            style.element.append(pf_xml)
        # 替换已有 tabs
        old_tabs = pf_xml.find(qn('w:tabs'))
        if old_tabs is not None:
            pf_xml.remove(old_tabs)
        pf_xml.append(tab_xml)

    # ── 设置打开文档时自动更新域 ──
    settings_el = doc.settings.element
    if settings_el.find(qn('w:updateFields')) is None:
        uf = OxmlElement('w:updateFields')
        uf.set(qn('w:val'), 'true')
        settings_el.append(uf)

    # ── TOC 域代码（begin + instrText + separate） ──
    p_field = doc.add_paragraph()
    r = p_field.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')
    r._r.append(fld_begin)

    r = p_field.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    r._r.append(instr)

    r = p_field.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fld_sep)

    # ── 预填充目录条目（兜底显示 + 页码占位） ──
    logger.info("Adding %d TOC entries as fallback", len(entries))
    style_map = {1: "toc 1", 2: "toc 2", 3: "toc 3"}
    for idx, (level, title) in enumerate(entries):
        if idx % 50 == 0 and idx > 0:
            logger.debug("TOC entry progress: %d/%d", idx, len(entries))
        ep = doc.add_paragraph(style=style_map.get(level, "toc 3"))
        ep.add_run(_clean(title))
        # 添加制表符和页码占位（Word 更新域后会替换）
        ep.add_run("\t")
        ep.add_run("__")

    # ── TOC 域代码（end） ──
    p_end = doc.add_paragraph()
    r = p_end.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fld_end)

    # ── 分节符（目录独立一节） ──
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2)
    new_section.left_margin = Cm(2.54)
    new_section.right_margin = Cm(2.54)

    logger.info("TOC generation completed")


# ==================== Word 导出 ====================

_DIAGRAM_KEYS = ("arch_diagram", "uml_diagram", "ui_diagrams")
_DIAGRAM_TITLES = {
    "arch_diagram": "系统架构图",
    "uml_diagram":  "UML 类图",
    "ui_diagrams":  "功能界面图",
}


def _add_diagram_chapters(doc, sections: list[GenerationSection], base_ch: int) -> None:
    """将已渲染的图表 PNG 插入 Word 文档，每个图表作为独立章节。"""
    from app.services.diagram_renderer import render_diagram_section

    ch = base_ch
    for key in _DIAGRAM_KEYS:
        sec = next((s for s in sections if s.section_key == key
                    and s.status == "completed" and s.content), None)
        if not sec:
            continue
        png = render_diagram_section(key, sec.content)
        if not png:
            continue

        cn = _CN_CHAPTER_NUMS[ch] if ch < len(_CN_CHAPTER_NUMS) else str(ch + 1)
        title = _DIAGRAM_TITLES.get(key, key)
        _add_heading(doc, f"{cn}、{title}", level=1)
        _add_diagram_to_doc(doc, png, key)
        doc.add_paragraph()  # 图片后空一行
        ch += 1


def export_manual_to_word(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出「文档鉴别材料」Word 文档（操作说明书）

    格式：宋体 12pt，固定行距 22pt，~31 行/页（≥30 行）
    如果 app.generate_diagrams=True，将 mermaid/html 代码块渲染为内嵌图片。
    """
    doc = _init_word_doc()
    render_diagrams = getattr(app, "generate_diagrams", False)

    manual_sections = [s for s in sections if s.section_key.startswith("manual_")]

    # 合并章节（功能详述上下 → 功能详述）
    chapters = _merge_manual_chapters(manual_sections)

    # 构建目录条目
    toc_entries = _build_toc_entries(chapters)

    # 封面（含分页符）
    _add_cover_page(doc, app, "使用说明书")

    # 目录
    _add_word_toc(doc, toc_entries)

    # 正文（图表内嵌到对应章节）
    for i, (title, content) in enumerate(chapters):
        cn = _CN_CHAPTER_NUMS[i] if i < len(_CN_CHAPTER_NUMS) else str(i + 1)
        _add_heading(doc, f"{cn}、{title}", level=1)
        counter = _HeadingCounter(i + 1)
        _add_manual_content(doc, content, counter=counter, skip_first_h1=title, render_diagrams=render_diagrams)

    return _save_doc(doc)


def export_source_code_to_word(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出「源程序鉴别材料」Word 文档（数据库设计 + 源程序代码）

    数据库设计：宋体 12pt，固定行距 22pt
    源程序代码：Times New Roman 10.5pt，固定行距 13pt，~53 行/页（≥50 行）
    """
    doc = _init_word_doc()

    db_sections = [s for s in sections if s.section_key.startswith("db_design")]
    code_sections = [s for s in sections if s.section_key.startswith("source_code_")]

    # 构建目录条目
    toc_entries: list[tuple[int, str]] = []
    chapter_num = 0
    if db_sections:
        chapter_num += 1
        cn = _CN_CHAPTER_NUMS[chapter_num - 1]
        toc_entries.append((1, f"{cn}、数据库设计"))
        for s in db_sections:
            if s.content:
                h2_count = 0
                h3_count = 0
                for h_level, h_title in _extract_headings(s.content):
                    if h_level == 1:
                        h2_count += 1
                        h3_count = 0
                        toc_entries.append((2, f"{chapter_num}.{h2_count}  {h_title}"))
                    elif h_level >= 2:
                        h3_count += 1
                        toc_entries.append((3, f"{chapter_num}.{h2_count}.{h3_count}  {h_title}"))
    if code_sections:
        chapter_num += 1
        cn = _CN_CHAPTER_NUMS[chapter_num - 1]
        toc_entries.append((1, f"{cn}、源程序代码"))

    # 封面（含分页符）
    _add_cover_page(doc, app, "源程序鉴别材料")

    # 目录
    _add_word_toc(doc, toc_entries)

    # 数据库设计
    ch = 0
    if db_sections:
        ch += 1
        cn = _CN_CHAPTER_NUMS[ch - 1]
        _add_heading(doc, f"{cn}、数据库设计", level=1)
        counter = _HeadingCounter(ch)
        for section in db_sections:
            if section.content:
                # 数据库设计一般不会有重复标题，传 None
                _add_manual_content(doc, section.content, counter=counter, skip_first_h1=None)

    # 源程序代码
    if code_sections:
        ch += 1
        if db_sections:
            doc.add_page_break()
        cn = _CN_CHAPTER_NUMS[ch - 1]
        _add_heading(doc, f"{cn}、源程序代码", level=1)
        for section in code_sections:
            if section.content:
                _add_source_code_content(doc, section.content)

    return _save_doc(doc)


def export_to_word(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出合并 Word 文档（向后兼容）"""
    doc = _init_word_doc()

    # 封面（含分页符）
    _add_cover_page(doc, app, "软件著作权申请材料")

    manual_sections = [s for s in sections if s.section_key.startswith("manual_")]
    db_sections = [s for s in sections if s.section_key.startswith("db_design")]
    code_sections = [s for s in sections if s.section_key.startswith("source_code_")]

    if manual_sections:
        doc.add_page_break()
        _add_heading(doc, "第一部分  操作说明书", level=1)
        for section in manual_sections:
            if section.content:
                _add_heading(doc, section.title, level=2)
                # 跳过内容中与 section 标题相同的第一个一级标题
                _add_manual_content(doc, section.content, base_heading_level=3, skip_first_h1=section.title)

    if db_sections:
        doc.add_page_break()
        _add_heading(doc, "第二部分  数据库设计", level=1)
        for section in db_sections:
            if section.content:
                _add_manual_content(doc, section.content, base_heading_level=2, skip_first_h1=None)

    if code_sections:
        doc.add_page_break()
        part_num = "第三部分" if db_sections else "第二部分"
        _add_heading(doc, f"{part_num}  源程序代码", level=1)
        for section in code_sections:
            if section.content:
                _add_source_code_content(doc, section.content)

    return _save_doc(doc)


# ==================== PDF 导出 ====================

_CN_FONT_SONG = "STSong-Light"  # 宋体（CID 内置）
_CN_FONT_HEI = None  # 黑体（运行时注册，可能为 None）


_fonts_registered = False


def _register_cn_fonts():
    global _CN_FONT_HEI, _fonts_registered
    if _fonts_registered:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    if _CN_FONT_SONG not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(_CN_FONT_SONG))

    # 尝试注册黑体 TTF，失败则回退到 STSong-Light
    if _CN_FONT_HEI is None:
        import os
        from reportlab.pdfbase.ttfonts import TTFont
        hei_paths = [
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Supplemental/SimHei.ttf",
            "/Library/Fonts/SimHei.ttf",
            "/usr/share/fonts/truetype/simhei/SimHei.ttf",
            "/usr/share/fonts/chinese/SimHei.ttf",
        ]
        for path in hei_paths:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont("Heiti", path, subfontIndex=0))
                    _CN_FONT_HEI = "Heiti"
                    break
                except Exception:
                    continue
        if _CN_FONT_HEI is None:
            _CN_FONT_HEI = _CN_FONT_SONG  # 回退
    _fonts_registered = True


def _init_pdf_styles():
    """初始化 PDF 样式集"""
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER

    _register_cn_fonts()
    styles = getSampleStyleSheet()

    return {
        "title": ParagraphStyle(
            "CnTitle", parent=styles["Title"],
            fontName=_CN_FONT_HEI, fontSize=22, leading=28, alignment=TA_CENTER,
        ),
        "h1": ParagraphStyle(
            "CnH1", parent=styles["Heading1"],
            fontName=_CN_FONT_HEI, fontSize=18, leading=24, spaceBefore=20, spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "CnH2", parent=styles["Heading2"],
            fontName=_CN_FONT_HEI, fontSize=15, leading=20, spaceBefore=14, spaceAfter=8,
        ),
        "h3": ParagraphStyle(
            "CnH3", parent=styles["Heading3"],
            fontName=_CN_FONT_HEI, fontSize=13, leading=18, spaceBefore=10, spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "CnBody", parent=styles["Normal"],
            fontName=_CN_FONT_SONG, fontSize=11, leading=18, spaceBefore=3, spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "CnBullet", parent=styles["Normal"],
            fontName=_CN_FONT_SONG, fontSize=11, leading=18, spaceBefore=3, spaceAfter=3,
            leftIndent=1.5 * cm, bulletIndent=0.5 * cm,
            bulletFontName=_CN_FONT_SONG, bulletFontSize=11,
        ),
        "code": ParagraphStyle(
            "CnCode", parent=styles["Code"],
            fontName=_CN_FONT_SONG, fontSize=8, leading=11,
            leftIndent=1 * cm, spaceBefore=4, spaceAfter=4,
        ),
        "code_line": ParagraphStyle(
            "CnCodeLine", parent=styles["Code"],
            fontName=_CN_FONT_SONG, fontSize=12, leading=22,
            spaceBefore=0, spaceAfter=0,
        ),
    }


def _build_pdf(story: list) -> io.BytesIO:
    """构建 PDF 文档"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2.54 * cm, rightMargin=2.54 * cm,
    )
    doc.build(story)
    buf.seek(0)
    return buf


class _TocDocTemplate:
    """支持自动目录的 PDF 文档模板（通过 multiBuild 两遍渲染）"""

    @staticmethod
    def build(story: list) -> io.BytesIO:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            topMargin=2 * cm, bottomMargin=2 * cm,
            leftMargin=2.54 * cm, rightMargin=2.54 * cm,
        )

        # 为 afterFlowable 注册 TOC 通知
        _original_after = doc.afterFlowable if hasattr(doc, 'afterFlowable') else lambda f: None

        def _after_flowable(flowable):
            _original_after(flowable)
            if hasattr(flowable, 'style'):
                name = flowable.style.name
                toc_map = {'CnH1': 0, 'CnH2': 1, 'CnH3': 2}
                if name in toc_map:
                    level = toc_map[name]
                    text = flowable.getPlainText()
                    doc.notify('TOCEntry', (level, text, doc.page))

        doc.afterFlowable = _after_flowable
        doc.multiBuild(story)
        buf.seek(0)
        return buf


def _build_pdf_with_toc(story: list, toc_insert_index: int = 4) -> io.BytesIO:
    """构建带目录的 PDF 文档

    Args:
        story: 已构建的 flowable 列表（含封面）
        toc_insert_index: 在 story 中插入目录的位置（默认 4，即封面 PageBreak 之后）
    """
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import Paragraph, PageBreak
    from reportlab.platypus.tableofcontents import TableOfContents

    _register_cn_fonts()

    # 目录标题
    toc_title = Paragraph("目  录", ParagraphStyle(
        'TocTitle', fontName=_CN_FONT_HEI, fontSize=18,
        leading=24, alignment=TA_CENTER, spaceBefore=40, spaceAfter=24,
    ))

    # 目录组件（紧凑样式，缩进为 0）
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle('TOCLevel0', fontName=_CN_FONT_HEI, fontSize=12, leading=20,
                       leftIndent=0, spaceBefore=4, spaceAfter=2),
        ParagraphStyle('TOCLevel1', fontName=_CN_FONT_SONG, fontSize=11, leading=18,
                       leftIndent=0, spaceBefore=2, spaceAfter=1),
        ParagraphStyle('TOCLevel2', fontName=_CN_FONT_SONG, fontSize=10, leading=16,
                       leftIndent=0, spaceBefore=1, spaceAfter=1),
    ]

    story.insert(toc_insert_index, toc_title)
    story.insert(toc_insert_index + 1, toc)
    story.insert(toc_insert_index + 2, PageBreak())

    return _TocDocTemplate.build(story)


_FONTCONFIG_CN_SUBST = """\
<?xml version="1.0"?>
<!DOCTYPE fontconfig SYSTEM "fonts.dtd">
<fontconfig>
  <!-- 继承系统字体配置 -->
  <include>/etc/fonts/fonts.conf</include>
  <!-- 将 Word 常用中文字体映射到服务器上已安装的 Noto/WQY 字体 -->
  <alias binding="same"><family>宋体</family>     <accept><family>Noto Serif CJK SC</family><family>WenQuanYi Micro Hei</family></accept></alias>
  <alias binding="same"><family>SimSun</family>   <accept><family>Noto Serif CJK SC</family><family>WenQuanYi Micro Hei</family></accept></alias>
  <alias binding="same"><family>黑体</family>     <accept><family>Noto Sans CJK SC</family><family>WenQuanYi Zen Hei</family></accept></alias>
  <alias binding="same"><family>SimHei</family>   <accept><family>Noto Sans CJK SC</family><family>WenQuanYi Zen Hei</family></accept></alias>
  <alias binding="same"><family>Microsoft YaHei</family><accept><family>Noto Sans CJK SC</family><family>WenQuanYi Zen Hei</family></accept></alias>
  <alias binding="same"><family>楷体</family>     <accept><family>Noto Serif CJK SC</family><family>WenQuanYi Micro Hei</family></accept></alias>
  <alias binding="same"><family>仿宋</family>     <accept><family>Noto Serif CJK SC</family><family>WenQuanYi Micro Hei</family></accept></alias>
</fontconfig>
"""


def _docx_to_pdf(docx_buf: io.BytesIO) -> io.BytesIO:
    """将 DOCX BytesIO 缓冲区转换为 PDF，使用 LibreOffice headless。

    服务器安装（Ubuntu/Debian）：
        apt-get install -y libreoffice --no-install-recommends fonts-noto-cjk fonts-wqy-zenhei
    """
    import subprocess
    import tempfile
    import os

    # 查找 libreoffice 可执行文件
    lo_bin = None
    for candidate in ("libreoffice", "soffice"):
        r = subprocess.run(["which", candidate], capture_output=True)
        if r.returncode == 0:
            lo_bin = r.stdout.decode().strip()
            break
    if not lo_bin:
        raise RuntimeError(
            "LibreOffice 未安装，无法转换 PDF。"
            "请在服务器上执行：apt-get install -y libreoffice --no-install-recommends fonts-noto-cjk fonts-wqy-zenhei"
        )

    docx_buf.seek(0)
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        profile_dir = os.path.join(tmpdir, "lo_profile")
        os.makedirs(profile_dir, exist_ok=True)

        with open(docx_path, "wb") as f:
            f.write(docx_buf.read())

        # 写入 fontconfig 中文字体替换规则
        fc_file = os.path.join(tmpdir, "fonts.conf")
        with open(fc_file, "w", encoding="utf-8") as f:
            f.write(_FONTCONFIG_CN_SUBST)

        env = os.environ.copy()
        env["FONTCONFIG_FILE"] = fc_file  # 覆盖 fontconfig 入口，注入替换规则

        result = subprocess.run(
            [
                lo_bin, "--headless", "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                docx_path,
            ],
            capture_output=True,
            timeout=300,
            env=env,
        )

        if result.returncode != 0:
            stderr = result.stderr.decode(errors="replace").strip()
            raise RuntimeError(f"LibreOffice 转换失败（code={result.returncode}）: {stderr[:500]}")

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice 未生成 PDF 文件，转换可能静默失败")

        with open(pdf_path, "rb") as f:
            return io.BytesIO(f.read())


def export_manual_to_pdf(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出「文档鉴别材料」PDF（操作说明书）

    先生成 Word 文档，再通过 Aspose.Words 转换为 PDF，排版更准确。
    """
    docx_buf = export_manual_to_word(app, sections)
    return _docx_to_pdf(docx_buf)


def export_source_code_to_pdf(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出「源程序鉴别材料」PDF（数据库设计 + 源程序代码）

    先生成 Word 文档，再通过 Aspose.Words 转换为 PDF，排版更准确。
    """
    docx_buf = export_source_code_to_word(app, sections)
    return _docx_to_pdf(docx_buf)


def export_to_pdf(
    app: Application, sections: list[GenerationSection]
) -> io.BytesIO:
    """导出合并 PDF 文档（向后兼容）

    先生成 Word 文档，再通过 Aspose.Words 转换为 PDF。
    """
    docx_buf = export_to_word(app, sections)
    return _docx_to_pdf(docx_buf)


# ==================== PDF 辅助 ====================

def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _add_pdf_table(story: list, table_lines: list[str], s_body) -> None:
    """将 Markdown 表格行渲染为 PDF 表格"""
    from reportlab.platypus import Table, Paragraph
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm

    rows_data: list[list[str]] = []
    for line in table_lines:
        if _is_separator_row(line):
            continue
        rows_data.append(_parse_table_row(line))

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)

    # 按内容长度比例分配列宽
    available_width = A4[0] - 2 * 2.54 * cm
    col_max_lens = [0] * num_cols
    for row_data in rows_data:
        for j in range(min(len(row_data), num_cols)):
            col_max_lens[j] = max(col_max_lens[j], len(row_data[j]))
    total_len = sum(col_max_lens) or 1
    col_widths = [available_width * l / total_len for l in col_max_lens]
    # 确保每列最小宽度
    min_col = 1.5 * cm
    for i in range(num_cols):
        if col_widths[i] < min_col:
            col_widths[i] = min_col

    # 构建表格数据（每个单元格是 Paragraph）
    table_data = []
    for i, row_data in enumerate(rows_data):
        row = []
        for j in range(num_cols):
            cell_text = row_data[j] if j < len(row_data) else ""
            if i == 0:
                row.append(Paragraph(f"<b>{_esc(cell_text)}</b>", s_body))
            else:
                row.append(Paragraph(_inline_format(_esc(cell_text)), s_body))
        table_data.append(row)

    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.Color(0.6, 0.6, 0.6)),
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.93, 0.93, 0.93)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ])
    story.append(t)


def _md_to_story(
    content: str, story: list, s_body, s_h3, s_bullet, s_code,
    *, s_h2=None, counter: _HeadingCounter | None = None,
) -> None:
    """简易 Markdown → reportlab flowables

    Args:
        s_h2: Heading 2 样式（仅在 counter 模式下使用）
        counter: 提供时，# → h2 (编号 N.1)，## → h3 (编号 N.1.1)
    """
    from reportlab.platypus import Spacer, Preformatted, Paragraph

    lines = content.split("\n")
    code_block = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    for line in lines:
        stripped = line.strip()

        # 如果正在收集表格行
        if table_lines:
            if stripped.startswith("|"):
                table_lines.append(stripped)
                continue
            else:
                _add_pdf_table(story, table_lines, s_body)
                table_lines = []

        if stripped.startswith("```"):
            if code_block:
                code_text = "\n".join(code_lines)
                story.append(Preformatted(code_text, s_code))
                code_lines = []
                code_block = False
            else:
                code_block = True
            continue

        if code_block:
            code_lines.append(line)
            continue

        if stripped.startswith("#"):
            hashes = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped.lstrip("#").strip()
            if title_text:
                if counter:
                    if hashes == 1:
                        numbered = counter.next_h2(title_text)
                        story.append(Paragraph(_esc(numbered), s_h2 or s_h3))
                    elif hashes >= 2:
                        numbered = counter.next_h3(title_text)
                        story.append(Paragraph(_esc(numbered), s_h3))
                else:
                    story.append(Paragraph(_esc(title_text), s_h3))
                continue

        # 表格开始
        if stripped.startswith("|"):
            table_lines.append(stripped)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            story.append(Paragraph(f"• {_inline_format(_esc(text))}", s_bullet))
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            story.append(Paragraph(_inline_format(_esc(text)), s_bullet))
        elif stripped == "":
            story.append(Spacer(1, 6))
        else:
            para_text = _inline_format(_esc(stripped))
            story.append(Paragraph(para_text, s_body))

    # 未关闭的表格
    if table_lines:
        _add_pdf_table(story, table_lines, s_body)

    if code_lines:
        code_text = "\n".join(code_lines)
        story.append(Preformatted(code_text, s_code))


def _source_code_to_story(content: str, story: list, s_code_line) -> None:
    """源代码 → reportlab flowables，分块输出提高性能"""
    from reportlab.platypus import Preformatted

    # 去除 markdown 代码块标记
    text = content.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        cleaned = []
        in_code = False
        for line in lines:
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            cleaned.append(line)
        text = "\n".join(cleaned) if cleaned else text

    # 按 200 行分块，用 Preformatted 批量输出（比逐行 Paragraph 快 10x+）
    all_lines = text.split("\n")
    chunk_size = 200
    for i in range(0, len(all_lines), chunk_size):
        chunk = "\n".join(all_lines[i:i + chunk_size])
        story.append(Preformatted(chunk, s_code_line))


def _inline_format(text: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
